# event_summary_app.py
import streamlit as st
from datetime import datetime, timedelta
import pandas as pd
import io
import re
from dateutil import parser as dateparser
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX


st.set_page_config(page_title="文件日期篩選摘要器", layout="wide")
st.title("📄 Word 文件 — 找出含指定日期的欄位並產生摘要檔")

st.markdown(
    """
上傳 `.docx`（可含表格或段落），選擇：  
- 使用「前一個工作日（前一個營業日）」或輸入**指定日期**，  
- 選擇要比對的欄位名稱（若 Word 內為表格會列出欄位供選），  
程式會找出在該欄位內含有目標日期（或日期字串）的列，並產生摘要檔下載。
"""
)

# -----------------------
# Helpers
# -----------------------
def prev_business_day(ref_date=None):
    # 計算前一個工作日（排除週六日）
    if ref_date is None:
        ref = datetime.now()
    else:
        ref = ref_date
    one = timedelta(days=1)
    d = ref - one
    while d.weekday() >= 5: # 5=Sat,6=Sun
        d -= one
    return d.date()

date_regex = re.compile(
    r'(?:(?:\d{4}[/-]\d{1,2}[/-]\d{1,2})|(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(?:\d{1,2}[\u4e00-\u9fff]{1}\d{1,2}[\u4e00-\u9fff]{0,1}\d{0,2}))'
)

def extract_tables_to_dfs(doc):
    """從 docx.Document 取出所有表格，回傳 list of DataFrame（盡量轉成 header）"""
    dfs = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        # 若表格至少有 2 列，嘗試將第一列當 header
        if len(rows) >= 2:
            header = rows[0]
            data = rows[1:]
            try:
                df = pd.DataFrame(data, columns=header)
            except Exception:
                # fallback: numeric columns
                df = pd.DataFrame(data)
        else:
            df = pd.DataFrame(rows)
        dfs.append(df)
    return dfs

def find_date_like_in_text(text):
    """嘗試用 regex 找日期字串並解析成 date，回傳 list of date objects（或空）"""
    found = []
    for m in date_regex.findall(text):
        s = m
        try:
            dt = dateparser.parse(s, dayfirst=False, fuzzy=True)
            if dt:
                found.append(dt.date())
        except Exception:
            continue
    return found

def filter_df_by_date_in_column(df, column, target_date):
    """
    在 DataFrame 的 column 欄內查找含有 target_date 的列。
    我們檢查：
      1) cell 文字中是否包含直接格式化日期（yyyy-mm-dd / yyyy/mm/dd / dd/mm/yyyy / 中文日期）
      2) 若 cell 本身是 datetime-like（pandas 解析過），也比對 date
    """
    if column not in df.columns:
        return pd.DataFrame()
    matches = []
    for idx, cell in df[column].fillna("").items():
        text = str(cell)
        # 1) 直接字串比對 target_date 的多種格式
        td_strs = [
            target_date.strftime("%Y-%m-%d"),
            target_date.strftime("%Y/%m/%d"),
            target_date.strftime("%Y%m%d"),
            target_date.strftime("%m/%d"),
            target_date.strftime("%#m/%#d/%Y") if hasattr(target_date, 'strftime') else ""
        ]
        # also Chinese style e.g. 2025年10月22日 and day/month permutations
        td_chinese = f"{target_date.year}年{target_date.month}月{target_date.day}日"
        found = False
        if any(s in text for s in td_strs) or td_chinese in text:
            found = True
        else:
            # 2) 用 regex 試著 parse cell 中的日期，再比對
            parsed = find_date_like_in_text(text)
            if parsed and any(d == target_date for d in parsed):
                found = True
        if found:
            matches.append((idx, cell))
    if not matches:
        return pd.DataFrame()
    # 保留整列
    return df.loc[[idx for idx, _ in matches]]

def export_to_word(data, target_date_str, filename="摘要輸出.docx"):
    doc = Document()
    doc.add_heading("搜尋摘要結果", level=1)
    doc.add_paragraph(f"搜尋日期關鍵字：{target_date_str}")
    doc.add_paragraph(" ")

    for idx, row in enumerate(data, start=1):
        p = doc.add_paragraph(f"{idx}. ")
        run = p.add_run(row)
        # 反白日期
        if target_date_str in row:
            # 這裡的 target_date_str 是例如 "2025-10-22"
            start = row.find(target_date_str)
            before = row[:start]
            date_part = row[start:start+len(target_date_str)]
            after = row[start+len(target_date_str):]

            p.clear()  # 清掉原本 run
            p.add_run(before)
            run_date = p.add_run(date_part)
            run_date.font.highlight_color = WD_COLOR_INDEX.YELLOW  # 反白日期
            p.add_run(after)

    doc.save(filename)
    return filename

# -----------------------
# UI: 上傳與選項
# -----------------------
uploaded_file = st.file_uploader("上傳 Word (.docx) 檔案", type=["docx"])
st.write("（檔案內容只在本次執行中處理，不會上傳到任何外部伺服器）")

col1, col2 = st.columns([2, 1])

with col1:
    date_mode = st.radio("選擇目標日期：", ("前一個工作日", "輸入指定日期 (YYYY-MM-DD)"))
    if date_mode == "輸入指定日期 (YYYY-MM-DD)":
        user_date_str = st.text_input("指定日期 (例: 2025-10-22)", value="")
        try:
            user_date = dateparser.parse(user_date_str).date() if user_date_str.strip() else None
        except Exception:
            user_date = None
    else:
        user_date = None

with col2:
    st.write("進階選項")
    prefer_table = st.checkbox("優先從表格欄位篩選 (若檔案含表格則列出欄位)", value=True)
    download_format = st.selectbox("下載摘要格式", ["CSV", "純文字 (TXT)", "Word (.docx)"])

# -----------------------
# 處理檔案
# -----------------------
if uploaded_file is not None:
    try:
        doc = docx.Document(uploaded_file)
    except Exception as e:
        st.error(f"無法讀取 Word 檔：{e}")
        st.stop()

    # 決定 target_date
    if date_mode == "前一個工作日":
        target_date = prev_business_day()
    else:
        if user_date is None:
            st.warning("請輸入有效的指定日期（YYYY-MM-DD）。")
            st.stop()
        target_date = user_date

    st.info(f"將比對的目標日期： {target_date.isoformat()}")

    # 嘗試從表格抓 dataframe
    dfs = extract_tables_to_dfs(doc)

    result_rows = []
    result_sources = [] # 說明是從哪個表格/段落找到

    if dfs and prefer_table:
        st.write(f"偵測到 {len(dfs)} 個表格，正在掃描表格欄位...")
        # 列出所有欄位供使用者選擇（合併所有表格欄位）
        all_cols = set()
        for df in dfs:
            all_cols.update(list(df.columns))
        all_cols = [c for c in all_cols if str(c).strip() != ""]
        if all_cols:
            chosen_cols = st.multiselect("選擇要比對的欄位（表格欄位）", options=all_cols, default=all_cols[:2])
        else:
            chosen_cols = []

        if chosen_cols:
            for i, df in enumerate(dfs):
                # 盡量讓 df 的列值為 str
                df = df.astype(str)
                for col in chosen_cols:
                    filtered = filter_df_by_date_in_column(df, col, target_date)
                    if not filtered.empty:
                        # 記錄並標註來源
                        filtered = filtered.copy()
                        filtered["_source_table"] = f"table_{i+1}"
                        filtered["_matched_column"] = col
                        result_rows.append(filtered)
        else:
            st.write("未選擇任何欄位 — 跳過表格欄位篩選。")

    # 若沒有結果或使用者不偏好表格，從段落中找（包含標籤：例如 '日期:'）
    if not result_rows:
        st.write("從段落中搜尋含有目標日期的文字（包含 '日期:','Date:' 等）...")
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        para_matches = []
        for i, txt in enumerate(paragraphs):
            parsed = find_date_like_in_text(txt)
            if parsed and any(d == target_date for d in parsed):
                para_matches.append((i, txt))
            else:
                # 或者包含文字型式 '前一個工作日' 或精確字串比對
                if target_date.isoformat() in txt or f"{target_date.year}年{target_date.month}月{target_date.day}日" in txt:
                    para_matches.append((i, txt))
        if para_matches:
            dfp = pd.DataFrame(para_matches, columns=["para_index", "text"])
            dfp["_source_table"] = "paragraphs"
            dfp["_matched_column"] = "text"
            result_rows.append(dfp)

    # 合併並顯示結果
    if result_rows:
        final = pd.concat(result_rows, ignore_index=True, sort=False)
        st.subheader("找到的結果範例（前 50 列）")
        st.dataframe(final.head(50))
        st.write(f"共找到 {len(final)} 筆符合目標日期 ({target_date}) 的項目。")

        # 產生下載檔案
        # -------- CSV 匯出 --------
        if download_format == "CSV":
            towrite = io.StringIO()
            final.to_csv(towrite, index=False, encoding="utf-8-sig")
            st.download_button(
                "下載 CSV 檔（UTF-8-SIG）",
                data=towrite.getvalue().encode("utf-8-sig"),
                file_name=f"summary_{target_date}.csv",
                mime="text/csv"
            )

        # -------- TXT 匯出 --------
        elif download_format == "純文字 (TXT)":
            txt_buf = io.StringIO()
            for i, row in final.iterrows():
                txt_buf.write(f"來源: {row.get('_table','')}, 欄位: {row.get('_matched_col','')}\n")
                for c in final.columns:
                    if not c.startswith("_"):
                        txt_buf.write(f"{c}: {row.get(c,'')}\n")
                txt_buf.write("----\n")
            st.download_button(
                "下載 TXT 檔",
                data=txt_buf.getvalue().encode("utf-8"),
                file_name=f"summary_{target_date}.txt",
                mime="text/plain"
            )

        # -------- Word 匯出 --------
        elif download_format == "Word (.docx)":

             st.write(f"進入WORD 格式處理")
             #輸出到 Word 並反白
             export_to_word(final, target_date)
            
             st.write(f"WORD 格式 完成")
             
             #取出 反白 Word 檔內容
             out_doc = docx.Document(“摘要輸出.docx”)
             word_stream = io.BytesIO()
             out_doc.save(word_stream)
             word_stream.seek(0)
             st.download_button(
                "下載 Word 摘要檔",
                data= word_stream,
                file_name=f"summary_{target_date}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
             )

            

    else:
        st.warning("沒有找到符合條件的項目。請確認：\n- Word 是否含有表格，或相關段落中是否有日期字串。\n- 若檔案使用特殊日期格式（例如中文全形空白或非標準符號），可手動輸入精確日期字串作為比對條件。")
