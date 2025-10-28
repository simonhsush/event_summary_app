# app.py
import streamlit as st
from datetime import datetime, timedelta
import pandas as pd
import io
import re
from dateutil import parser as dateparser
import docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

st.set_page_config(page_title="文件日期篩選摘要器", layout="wide")
st.title("📄 Word 文件 — 找出含指定日期的欄位並產生摘要檔")

st.markdown(
    """
上傳 `.docx`（可含表格或段落），選擇：
- 使用「前一個工作日（前一個營業日）」或輸入**指定日期**，
- 選擇要比對的欄位名稱（若 Word 內為表格會列出欄位供選），
程式會找出在該欄位內含有目標日期（或日期字串）的列，並產生摘要檔下載。
"""
)

# -----------------------
# 日期處理邏輯（保留原樣）
# -----------------------
def prev_business_day(ref_date=None):
    if ref_date is None:
        ref = datetime.now()
    else:
        ref = ref_date
    one = timedelta(days=1)
    d = ref - one
    while d.weekday() >= 5:
        d -= one
    return d.date()

date_regex = re.compile(
    r'(?:(?:\d{4}[/-]\d{1,2}[/-]\d{1,2})|(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4})|(?:\d{1,2}[\u4e00-\u9fff]{1}\d{1,2}[\u4e00-\u9fff]{0,1}\d{0,2}))'
)

def extract_tables_to_dfs(doc):
    dfs = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        if len(rows) >= 2:
            header = rows[0]
            data = rows[1:]
            try:
                df = pd.DataFrame(data, columns=header)
            except Exception:
                df = pd.DataFrame(data)
        else:
            df = pd.DataFrame(rows)
        dfs.append(df)
    return dfs

def find_date_like_in_text(text):
    found = []
    for m in date_regex.findall(text):
        s = m
        try:
            dt = dateparser.parse(s, dayfirst=False, fuzzy=True)
            if dt:
                found.append(dt.date())
        except Exception:
            continue
    return found

# -----------------------
# 🔹 強化日期比對邏輯（支援 10/23、114/10/23、10_23、10月23日）
# -----------------------
def filter_df_by_date_in_column(df, column, target_date):
    if column not in df.columns:
        return pd.DataFrame()
    matches = []
    for idx, cell in df[column].fillna("").items():
        text = str(cell)

        # 🔹 將底線與全形符號正規化
        normalized_text = (
            text
            .replace("／", "/")
            .replace("－", "-")
            .replace("_", "/")
            .replace("．", ".")
        )

        # 🔹 原邏輯 + 擴充格式
        td_strs = [
            target_date.strftime("%Y-%m-%d"),
            target_date.strftime("%Y/%m/%d"),
            target_date.strftime("%Y%m%d"),
            target_date.strftime("%m/%d"),
            f"{target_date.month}/{target_date.day}", # ← 無年份
        ]

        td_chinese = f"{target_date.year}年{target_date.month}月{target_date.day}日"
        td_chinese_short = f"{target_date.month}月{target_date.day}日" # ← 中文短格式
        td_roc = f"{target_date.year - 1911}/{target_date.month}/{target_date.day}" # ← 民國年格式
        td_roc_no_year = f"{target_date.month}/{target_date.day}" # ← 民國年省略情況

        found = False
        # 🔹 改在 normalized_text 裡找
        if any(s in normalized_text for s in td_strs) or \
           td_chinese in normalized_text or \
           td_chinese_short in normalized_text or \
           td_roc in normalized_text or \
           td_roc_no_year in normalized_text:
            found = True
        else:
            parsed = find_date_like_in_text(normalized_text)
            if parsed and any(d == target_date for d in parsed):
                found = True

        if found:
            matches.append((idx, cell))

    if not matches:
        return pd.DataFrame()
    return df.loc[[idx for idx, _ in matches]]


# -----------------------
# 匯出 Word — 從日期後取指定字數
# -----------------------
def export_to_word(data, target_date_str, num_chars, filename="摘要輸出.docx"):
    doc = Document()
    doc.add_heading("搜尋摘要結果", level=1)
    doc.add_paragraph(f"搜尋日期關鍵字：{target_date_str}")
    doc.add_paragraph(" ")

    if "text" not in data.columns:
        data = data.copy()
        data["text"] = data.apply(lambda r: " ".join([str(x) for x in r.values if pd.notna(x)]), axis=1)

    for idx, row in enumerate(data["text"].astype(str).tolist(), start=1):
        text = row
        start_idx = text.find(target_date_str)
        if start_idx == -1:
            try:
                parts = re.findall(r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})', target_date_str)
                if parts:
                    y, m, d = parts[0]
                    chinese = f"{int(y)}年{int(m)}月{int(d)}日"
                    start_idx = text.find(chinese)
                    target_for_slice = chinese
                else:
                    target_for_slice = target_date_str
            except Exception:
                target_for_slice = target_date_str
        else:
            target_for_slice = target_date_str

        if start_idx != -1:
            slice_start = start_idx
            slice_end = min(len(text), slice_start + len(target_for_slice) + num_chars)
            snippet = text[slice_start:slice_end]
        else:
            snippet = ""

        p = doc.add_paragraph(f"{idx}. ")
        if snippet:
            run = p.add_run(snippet)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            p.add_run("(未找到可擷取的段落)")

    doc.save(filename)
    return filename


# -----------------------
# Streamlit UI
# -----------------------
uploaded_file = st.file_uploader("上傳 Word (.docx) 檔案", type=["docx"])
st.write("（檔案內容只在本次執行中處理，不會上傳到任何外部伺服器）")

num_chars = st.number_input("請輸入要擷取的字數（從日期字串之後開始計算，預設20）", min_value=1, max_value=1000, value=20)

col1, col2 = st.columns([2, 1])
with col1:
    date_mode = st.radio("選擇目標日期：", ("前一個工作日", "輸入指定日期 (YYYY-MM-DD)"))
    if date_mode == "輸入指定日期 (YYYY-MM-DD)":
        user_date_str = st.text_input("指定日期 (例: 2025-10-22)", value="")
        try:
            user_date = dateparser.parse(user_date_str).date() if user_date_str.strip() else None
        except Exception:
            user_date = None
    else:
        user_date = None

with col2:
    st.write("進階選項")
    prefer_table = st.checkbox("優先從表格欄位篩選 (若檔案含表格則列出欄位)", value=True)
    download_format = st.selectbox("下載摘要格式", ["CSV", "純文字 (TXT)", "Word (.docx)"])

# -----------------------
# 主流程
# -----------------------
if uploaded_file is not None:
    try:
        doc = docx.Document(uploaded_file)
    except Exception as e:
        st.error(f"無法讀取 Word 檔：{e}")
        st.stop()

    if date_mode == "前一個工作日":
        target_date = prev_business_day()
        target_date_str = target_date.isoformat()
    else:
        if user_date is None:
            st.warning("請輸入有效的指定日期（YYYY-MM-DD）。")
            st.stop()
        target_date = user_date
        target_date_str = str(target_date)

    st.info(f"將比對的目標日期： {target_date_str}")

    dfs = extract_tables_to_dfs(doc)
    result_rows = []

    if dfs and prefer_table:
    st.write(f"偵測到 {len(dfs)} 個表格，以下為各表格的預覽與比對設定：")

    result_rows = []

    for i, df in enumerate(dfs, start=1):
        with st.expander(f"📋 表格 {i} 預覽", expanded=True):
            df = df.astype(str)

            st.dataframe(df.head(50), use_container_width=True)

            # 自動抓出欄位名稱
            table_cols = [c for c in df.columns if str(c).strip() != ""]
            if not table_cols:
                st.warning("此表格沒有可辨識的欄位名稱（可能缺少標題列）。")
                continue

            # 預設所有欄位都勾選
            chosen_cols = st.multiselect(
                f"表格 {i} — 選擇要比對的欄位",
                options=table_cols,
                default=table_cols
            )

            # 若使用者沒有選任何欄位就跳過
            if not chosen_cols:
                continue

            # 執行比對
            for col in chosen_cols:
                filtered = filter_df_by_date_in_column(df, col, target_date)
                if not filtered.empty:
                    filtered = filtered.copy()
                    snippets = []
                    for _, r in filtered.iterrows():
                        cell_text = str(r[col])
                        td_candidates = [
                            target_date.strftime("%Y-%m-%d"),
                            target_date.strftime("%Y/%m/%d"),
                            f"{target_date.year}年{target_date.month}月{target_date.day}日"
                        ]
                        start_idx = -1
                        chosen_td = None
                        for td in td_candidates:
                            if td in cell_text:
                                start_idx = cell_text.find(td)
                                chosen_td = td
                                break
                        if start_idx != -1:
                            end_idx = min(len(cell_text), start_idx + len(chosen_td) + num_chars)
                            snippet = cell_text[start_idx:end_idx]
                        else:
                            snippet = ""
                        snippets.append(snippet)
                    filtered = filtered.reset_index(drop=True)
                    filtered["text"] = snippets
                    filtered["_source_table"] = f"table_{i}"
                    filtered["_matched_column"] = col
                    filtered = filtered[filtered["text"].str.strip() != ""]
                    if not filtered.empty:
                        result_rows.append(filtered)

    if not result_rows:
        st.write("從段落中搜尋含有目標日期的文字...")
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        para_matches = []

        for i, txt in enumerate(paragraphs):
            parsed = find_date_like_in_text(txt)
            if parsed and any(d == target_date for d in parsed):
                td_candidates = [
                    target_date.strftime("%Y-%m-%d"),
                    target_date.strftime("%Y/%m/%d"),
                    f"{target_date.year}年{target_date.month}月{target_date.day}日"
                ]
                start_idx = -1
                chosen_td = None
                for td in td_candidates:
                    if td in txt:
                        start_idx = txt.find(td)
                        chosen_td = td
                        break
                if start_idx != -1:
                    end_idx = min(len(txt), start_idx + len(chosen_td) + num_chars)
                    snippet = txt[start_idx:end_idx]
                    para_matches.append((i, snippet))

        if para_matches:
            dfp = pd.DataFrame(para_matches, columns=["para_index", "text"])
            dfp["_source_table"] = "paragraphs"
            dfp["_matched_column"] = "text"
            result_rows.append(dfp)

    if result_rows:
        final = pd.concat(result_rows, ignore_index=True, sort=False)
        final = final[final["text"].astype(str).str.strip() != ""].reset_index(drop=True)

        st.subheader("找到的結果範例（前 200 列）")
        st.dataframe(final.head(200))
        st.write(f"共找到 {len(final)} 筆符合目標日期 ({target_date_str}) 的項目。")

        if download_format == "CSV":
            towrite = io.StringIO()
            final.to_csv(towrite, index=False, encoding="utf-8-sig")
            st.download_button(
                "下載 CSV 檔（UTF-8-SIG）",
                data=towrite.getvalue().encode("utf-8-sig"),
                file_name=f"summary_{target_date_str}.csv",
                mime="text/csv"
            )
        elif download_format == "純文字 (TXT)":
            txt_buf = io.StringIO()
            for i, row in final.iterrows():
                txt_buf.write(f"來源: {row.get('_source_table','')}, 欄位: {row.get('_matched_column','')}\n")
                txt_buf.write(f"{row.get('text','')}\n")
                txt_buf.write("----\n")
            st.download_button(
                "下載 TXT 檔",
                data=txt_buf.getvalue().encode("utf-8"),
                file_name=f"summary_{target_date_str}.txt",
                mime="text/plain"
            )
        elif download_format == "Word (.docx)":
            export_to_word(final, target_date_str, num_chars, filename="日期段落摘要.docx")
            out_doc = docx.Document("日期段落摘要.docx")
            word_stream = io.BytesIO()
            out_doc.save(word_stream)
            word_stream.seek(0)
            st.download_button(
                "下載 Word 摘要檔",
                data=word_stream,
                file_name=f"summary_{target_date_str}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("沒有找到符合條件的項目。請確認：\n- Word 是否含有表格或段落中是否有日期字串。\n- 若日期格式特殊，可嘗試手動輸入精確日期字串作為比對條件。")



